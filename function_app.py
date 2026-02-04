import azure.functions as func
import logging
import os
import json
import requests
import ollama 
import random
import time
import fitz  # PyMuPDF
from docx import Document
from playwright.sync_api import sync_playwright
from dotenv import load_dotenv
from pathlib import Path
from datetime import datetime, timedelta

load_dotenv()
app = func.FunctionApp()

# --- CONFIGURATION ---
AUTO_SUBMIT = True
RESUME_PATH = "resume.pdf"
LOG_FILE = Path("applied_jobs.json")

BLACKLIST = ["Sales", "Account Executive", "Business Development", "SDR", "BDR", "Account Manager", "Customer Success", "Recruiter", "Marketing", "Head of"]
WHITELIST = ["Cyber", "Security", "AI", "Artificial Intelligence", "Analyst", "Engineer", "Developer", "Architect", "SOC", "Penetration", "Red Team", "Blue Team", "Automation", "Threat", "Vulnerability", "GRC", "Product", "Intern"]

def is_valid_role(title):
    for bad_word in BLACKLIST:
        if bad_word.lower() in title.lower():
            logging.info(f"⛔ Skipped Blacklisted: {title}")
            return False
    for good_word in WHITELIST:
        if good_word.lower() in title.lower():
            return True
    return False

def cleanup_old_resumes():
    output_dir = Path("tailored_resumes")
    if not output_dir.exists(): return
    now = datetime.now()
    cutoff = now - timedelta(hours=24)
    for file in output_dir.glob("*.docx"):
        if datetime.fromtimestamp(file.stat().st_mtime) < cutoff:
            file.unlink()

def has_already_applied(job_id):
    if not LOG_FILE.exists(): return False
    try:
        with open(LOG_FILE, "r") as f:
            applied_ids = json.load(f)
            return job_id in applied_ids
    except: return False

def log_applied_job(job_id):
    applied_ids = []
    if LOG_FILE.exists():
        try:
            with open(LOG_FILE, "r") as f: applied_ids = json.load(f)
        except: pass
    if job_id not in applied_ids:
        applied_ids.append(job_id)
        with open(LOG_FILE, "w") as f: json.dump(applied_ids, f, indent=4)

def extract_resume_text(pdf_path):
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc: text += page.get_text()
    except Exception as e: logging.error(f"PDF Error: {e}")
    return text

def create_tailored_resume(job_title, job_description):
    base_text = extract_resume_text(RESUME_PATH)
    logging.info(f"✍️ Forging tailored resume for: {job_title}")
    prompt = (
        f"Sam is a Senior Cyber student. Projects: Archangel, SEET Paper, LDI Connect GRC.\n"
        f"Job: {job_title}\n"
        f"Description: {job_description[:2000]}\n"
        "Task: Write a 3-sentence summary highlighting his AI and Cyber skills for this role."
    )
    response = ollama.chat(model='llama3.1', messages=[{'role': 'user', 'content': prompt}])
    new_summary = response['message']['content']
    doc = Document()
    doc.add_heading(f'Sam Oakes - {job_title}', 0)
    doc.add_paragraph(new_summary)
    output_dir = Path("tailored_resumes")
    output_dir.mkdir(exist_ok=True)
    clean_title = "".join([c for c in job_title if c.isalnum() or c==' ']).replace(' ', '_')
    file_path = output_dir / f"{clean_title}_Resume.docx"
    doc.save(str(file_path))
    return str(file_path.absolute())

def analyze_job_locally(job_title):
    logging.info(f"🧠 Llama 3.1 is analyzing: {job_title}")
    response = ollama.chat(model='llama3.1', messages=[
        {'role': 'system', 'content': "You are a career agent for Sam. Write a 1-sentence recruiter hook."}, 
        {'role': 'user', 'content': f"Job Title: {job_title}"},
    ])
    return response['message']['content']

def handle_easy_apply(page, tailored_resume_path):
    try:
        # 1. Clear overlays
        try:
            if page.is_visible(".msg-overlay-bubble-header__control--close-btn"):
                page.click(".msg-overlay-bubble-header__control--close-btn")
        except: pass

        # 2. Hunt for the button
        apply_btn = None
        selectors = [
            "button.jobs-apply-button", 
            "button:has-text('Easy Apply')",
            ".jobs-apply-button--top-card button"
        ]
        
        for s in selectors:
            if page.is_visible(s):
                apply_btn = page.query_selector(s)
                break
        
        if not apply_btn:
            logging.warning("Apply button hidden.")
            return False

        # 3. The "Double-Tap" Logic
        logging.info("🖱️ Clicking Apply Button...")
        apply_btn.click(force=True)
        time.sleep(2)

        # Check if modal appeared. If not, click again.
        if not page.is_visible(".artdeco-modal"):
            logging.info("🔄 Modal didn't appear. Double-tapping...")
            apply_btn.click(force=True)
            time.sleep(2)

        # 4. Wait for ANY modal (using the generic class)
        try:
            # .artdeco-modal is the master class for all LinkedIn popups
            page.wait_for_selector(".artdeco-modal", timeout=10000)
            logging.info("✅ Modal Detected.")
        except:
            logging.warning("❌ No modal appeared even after double-tap.")
            page.screenshot(path="modal_fail_debug.png")
            return False

        # 5. Fill the form
        for _ in range(8): # Increased steps for longer forms
            time.sleep(1.5)
            
            # Upload Resume if asked
            file_input = page.query_selector("input[type='file']")
            if file_input: 
                logging.info(f"📤 Uploading: {tailored_resume_path}")
                try:
                    file_input.set_input_files(tailored_resume_path)
                except: pass # Sometimes it's already filled

            next_btn = page.query_selector("button:has-text('Next')")
            review_btn = page.query_selector("button:has-text('Review')")
            submit_btn = page.query_selector("button:has-text('Submit application')")

            if submit_btn:
                if AUTO_SUBMIT: 
                    submit_btn.click()
                    logging.info("✅ FINAL SUBMIT CLICKED!")
                else:
                    logging.info("🚨 Safe Mode: Ready to Submit.")
                break
            
            if review_btn: review_btn.click()
            elif next_btn: next_btn.click()
            else: 
                # If no buttons found, we might be stuck or done
                break
        return True
            
    except Exception as e:
        logging.error(f"Easy Apply Error: {e}")
        page.screenshot(path="error_debug.png")
        return False

@app.timer_trigger(schedule="0 */5 * * * *", arg_name="myTimer", run_on_startup=True)
def JobScraper(myTimer: func.TimerRequest) -> None:
    run_automation()

def run_automation():
    cleanup_old_resumes()
    user_data_dir = str(Path.cwd() / "chrome_profile")
    with sync_playwright() as p:
        context = p.chromium.launch_persistent_context(
            user_data_dir, headless=False, slow_mo=1000,
            args=["--disable-blink-features=AutomationControlled"]
        )
        page = context.pages[0] if context.pages else context.new_page()
        
        logging.info("🌍 Navigating to EASY APPLY + REMOTE Cybersecurity & AI roles...")
        page.goto("https://www.linkedin.com/jobs/search/?keywords=Cybersecurity%20OR%20AI&f_AL=true&f_WT=2&location=United%20States")
        
        try:
            page.wait_for_selector(".job-card-container", timeout=15000)
            job_cards = page.query_selector_all(".job-card-container")
            
            found_valid_job = False
            for card in job_cards[:5]: 
                job_id = card.get_attribute("data-job-id")
                
                if job_id and has_already_applied(job_id):
                    logging.info(f"⏭️ Skipping {job_id}: Already processed.")
                    continue
                
                card.click()
                time.sleep(3) 

                # Title Extraction
                title_selectors = ["h2.t-24", ".job-details-jobs-unified-top-card__job-title", "h1.t-24"]
                title = None
                for s in title_selectors:
                    el = page.query_selector(s)
                    if el: 
                        title = el.inner_text().strip()
                        if title: break
                
                if not title: continue
                if not is_valid_role(title): continue

                logging.info(f"✅ LOCK ON: {title}")
                
                # Desc Extraction
                desc_selectors = [".jobs-search__job-details--container", "#job-details", ".jobs-box__html-content"]
                desc = None
                for s in desc_selectors:
                    el = page.query_selector(s)
                    if el: 
                        desc = el.inner_text().strip()
                        if desc and len(desc) > 50: break 

                if desc:
                    new_resume = create_tailored_resume(title, desc)
                    pitch = analyze_job_locally(title)
                    success = handle_easy_apply(page, new_resume)
                    
                    if success:
                        log_applied_job(job_id)
                        webhook_url = os.getenv("DISCORD_WEBHOOK_URL")
                        if webhook_url:
                            requests.post(webhook_url, json={"content": f"🎯 **Applied (Remote):** {title}\n**AI Pitch:** {pitch}"})
                        found_valid_job = True
                        break 
            
            if not found_valid_job:
                logging.info("🏁 No valid new jobs found in this batch.")

        except Exception as e:
            logging.error(f"Automation failed: {e}")
        
        logging.info("Loop complete. Cooling down...")
        time.sleep(10)
        context.close()